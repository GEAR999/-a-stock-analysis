#!/usr/bin/env python3
"""将 AGENTS.md 转换为 Word 文档"""

import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def read_markdown(filepath):
    with open(filepath, 'r', encoding='utf-8') as f:
        return f.read()

def create_word_from_md(md_content, output_path):
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(10)
    
    lines = md_content.split('\n')
    
    for line in lines:
        # 标题处理
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        # 代码块
        elif line.startswith('```'):
            pass  # 跳过代码块标记
        # 列表项
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        # 表格行
        elif line.startswith('|') and '|' in line[1:]:
            # 简单表格处理
            cells = [c.strip() for c in line.split('|')[1:-1]]
            if cells and not all(c == '' or c == '---' for c in cells):
                # 检查是否需要创建新表格
                table = doc.add_table(rows=1, cols=len(cells))
                for i, cell in enumerate(cells):
                    table.rows[0].cells[i].text = cell
        # 空行
        elif line.strip() == '':
            doc.add_paragraph('')
        # 普通文本
        else:
            # 处理行内代码
            text = re.sub(r'`([^`]+)`', r'\1', line)
            if text.strip():
                doc.add_paragraph(text)
    
    doc.save(output_path)
    print(f"Word 文档已生成：{output_path}")

if __name__ == '__main__':
    md_content = read_markdown('/workspace/projects/AGENTS.md')
    create_word_from_md(md_content, '/workspace/projects/AGENTS.docx')
